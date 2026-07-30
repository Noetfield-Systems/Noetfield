#!/usr/bin/env python3
"""NOETFIELD_SFF_FINAL_RECONCILIATION_V1 — regenerate SFF long-form PDF/DOCX with TrustField boundary."""

from __future__ import annotations

import hashlib
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
SOURCE_PDF = Path.home() / "Downloads/SFF_Noetfield_Long_Form_For_Profit_FINAL.pdf"
OUT_DIR = ROOT / "public-interest/artifacts"
OUT_PDF = OUT_DIR / "SFF_Noetfield_Long_Form_For_Profit_RECONCILIATION_V1.pdf"
OUT_DOCX = OUT_DIR / "SFF_Noetfield_Long_Form_For_Profit_RECONCILIATION_V1.docx"

GENERAL_ACTIVITIES_REPLACEMENT = (
    "Noetfield's current work includes the live client-zero execution application "
    "and SourceA / Runway execution and evidence infrastructure. TrustField is a "
    "Noetfield Systems Inc. product whose synthetic regulated-workflow demonstrations "
    "provide a bounded validation context."
)

FUNDED_ACTIVITIES_REPLACEMENT = (
    "Evaluate selected authority and evidence patterns against synthetic TrustField "
    "workflows within the product's stated boundary."
)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_text(pdf_path: Path) -> str:
    reader = PdfReader(str(pdf_path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def apply_replacements(text: str) -> str:
    out = text

    out = re.sub(
        r"Noetfield Systems Inc\. is a Vancouver-based AI-native systems company that builds governed execution infrastructure and\s+specialized workflow products\.\s+Its current work includes the live Noetfield client-zero application, SourceA/Runway execution\s+infrastructure, and TrustField regulated-operations workflows designed to preserve human decision authority and inspectable\s+evidence\.",
        (
            "Noetfield Systems Inc. is a Vancouver-based AI-native systems company that builds governed execution infrastructure. "
            + GENERAL_ACTIVITIES_REPLACEMENT
        ),
        out,
        flags=re.DOTALL,
    )

    out = out.replace(
        "TrustField regulated-operations workflows.",
        "TrustField product validation context.",
    )
    out = re.sub(
        r"TrustField provides live, synthetic workflow demonstrations for case intake, analyst\s+review, human FILE / NO-FILE / ESCALATE decisions, submission tracking, and evidence closure\.\s+The product explicitly states that it",
        (
            "Within TrustField's own scope, synthetic workflow demonstrations cover case intake, analyst "
            "review, human FILE / NO-FILE / ESCALATE decisions, submission tracking, and evidence closure. "
            "TrustField explicitly states that it"
        ),
        out,
        flags=re.DOTALL,
    )

    out = out.replace(
        "Validate through TrustField: use regulated-operations workflows to test human decision preservation, evidence completeness,  exception routing, and bounded automation without custody, settlement, auto-filing, or certification claims.",
        FUNDED_ACTIVITIES_REPLACEMENT
        + " Testing stays within the product's boundary — without custody, settlement, auto-filing, or certification claims.",
    )

    out = out.replace(
        "specialized workflow products such as TrustField.",
        "specialized Noetfield workflow products and licensed execution infrastructure.",
    )
    out = re.sub(
        r"(Early revenue may\s+include limited implementation work, but the long-term model is reusable software and execution infrastructure rather than selling\s+hours\.)",
        r"\1 TrustField is excluded from Noetfield Systems Inc.'s revenue model unless a later written inter-entity agreement establishes a specific transaction.",
        out,
    )

    out = out.replace(
        "TrustField, a live compliance-workflow platform.",
        "TrustField, a Noetfield Systems Inc. product whose synthetic workflows demonstrate human decision gates and evidence boundaries within TrustField's own scope.",
    )
    out = out.replace(
        "validate those  controls through TrustField's regulated-operations workflows",
        "evaluate selected patterns against synthetic TrustField workflows within the product's stated boundary",
    )
    out = out.replace(
        "TrustField as a high- consequence vertical",
        "TrustField, within the product's own scope, as a high-consequence validation context",
    )
    out = out.replace("28 July 2026", "29 July 2026")
    out = out.replace("www.noetfield.com/proof/noetfield/", "www.noetfield.com/proof/claims-boundary-correction/")
    out = out.replace("/proof/noetfield/", "/proof/claims-boundary-correction/")

    for amount in ("180,000", "25,000", "450,000", "750,000"):
        if amount not in out:
            raise ValueError(f"budget amount missing after edit: {amount}")

    return out


def write_docx(text: str, path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for para in text.split("\n"):
        p = para.strip()
        if not p:
            doc.add_paragraph("")
            continue
        doc.add_paragraph(p)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_pdf(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    margin = inch
    y = height - margin
    line_height = 12
    max_width = width - 2 * margin

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = height - margin

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            y -= line_height
            if y < margin:
                new_page()
            continue
        words = line.split()
        buf = ""
        for word in words:
            trial = (buf + " " + word).strip()
            if c.stringWidth(trial, "Times-Roman", 10) <= max_width:
                buf = trial
            else:
                if y < margin:
                    new_page()
                c.setFont("Times-Roman", 10)
                c.drawString(margin, y, buf)
                y -= line_height
                buf = word
        if buf:
            if y < margin:
                new_page()
            c.setFont("Times-Roman", 10)
            c.drawString(margin, y, buf)
            y -= line_height
    c.save()


def main() -> int:
    if not SOURCE_PDF.is_file():
        raise SystemExit(f"missing source PDF: {SOURCE_PDF}")
    text = apply_replacements(extract_text(SOURCE_PDF))
    write_docx(text, OUT_DOCX)
    write_pdf(text, OUT_PDF)
    print(f"wrote {OUT_PDF.relative_to(ROOT)} sha256={sha256_file(OUT_PDF)}")
    print(f"wrote {OUT_DOCX.relative_to(ROOT)} sha256={sha256_file(OUT_DOCX)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
